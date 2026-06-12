# generate_docx.py
import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_document():
    doc = docx.Document()

    # Set professional styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("EcoTrack: Carbon-Aware MLOps Governance System\n")
    run.bold = True
    run.font.size = Pt(20)

    sub = title.add_run("Academic Internship End-Project Review Portfolio\n")
    sub.italic = True
    sub.font.size = Pt(14)

    doc.add_paragraph("="*60).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # High Level Overview
    doc.add_heading("1. High-Level Project Overview", level=1)
    doc.add_paragraph(
        "EcoTrack is a carbon-aware MLOps governance middleware that addresses the environmental cost of modern artificial intelligence pipelines. "
        "Traditional machine learning operations trigger automated retraining cycles as soon as data drift is detected, consuming significant grid electricity "
        "and water cooling resources. EcoTrack introduces a governance layer (the Eco-Gate) that monitors regional grid carbon intensity, computes water cooling loss, "
        "and calculates compliance tax liabilities, dynamically regulating retraining based on ambient grid health."
    )

    # Architectural Design
    doc.add_heading("2. Architectural Design", level=1)
    doc.add_paragraph(
        "The system is built on a decoupled, three-tier architecture ensuring complete separation of concerns:"
    )
    doc.add_paragraph("• Presentation Layer (Frontend): Designed in React/Vite, featuring animated entry portals, telemetry meters, and live model comparison metrics.")
    doc.add_paragraph("• Logic Layer (Backend): Built using FastAPI and Uvicorn. It performs model evaluation, calculates compliance ledgers, and runs tournament duels.")
    doc.add_paragraph("• Data Layer (Storage & Registry): Integrates MongoDB to maintain prediction logs and versioned model registration documents, coupled with local file system binary serialization (.pkl).")

    # Contributor Roles
    doc.add_heading("3. Contributor Roles & Team Responsibilities", level=1)
    doc.add_paragraph(
        "The project was executed by three team members, divided logically across the tiers:"
    )

    # Deepa Sahana
    p1 = doc.add_paragraph()
    r1 = p1.add_run("Deepa Sahana — Data Collection & Model Training\n")
    r1.bold = True
    p1.add_run(
        "• Curated base air quality datasets with feature parameters (PM2.5, PM10, NO2, CO, SO2, AQI) matching central regulation standards.\n"
        "• Designed and trained the baseline Decision Tree Regressor model.\n"
        "• Implemented skewed noise generators to simulate 2026 climate/pollution data drift."
    )

    # Dhaarani
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Dhaarani — API Development & MLOps Governance\n")
    r2.bold = True
    p2.add_run(
        "• Developed the FastAPI backend service supporting RESTful endpoints (/status, /inject, /retrain, /rollback, /reset).\n"
        "• Engineered the Eco-Gate rule engine computing carbon footprint taxes and water loss rates.\n"
        "• Integrated MongoDB registry collection and resolved the version rollback sync bug by aligning dataset truncation with weights hot-swapping."
    )

    # Janani
    p3 = doc.add_paragraph()
    r3 = p3.add_run("Janani — Frontend UI & Client Integration\n")
    r3.bold = True
    p3.add_run(
        "• Developed the Single Page Application using React and Vite.\n"
        "• Crafted the animated realistic glassmorphism landing page and responsive graphs using Recharts.\n"
        "• Programmed layout controls, time sliders, and dashboard-exit navigation systems."
    )

    doc.add_heading("4. MLOps Implementation Strategies", level=1)
    doc.add_paragraph("• Environmental Gating: Continuous monitoring of carbon intensity to delay dirty retraining cycles.")
    doc.add_paragraph("• Model Tournament: Champion vs. Challenger duels to validate accuracy improvement prior to promotion.")
    doc.add_paragraph("• Double-Alignment Rollback: Rolling back model registry documents automatically reverts data states to maintain consistency.")

    doc.save("project_overview.docx")
    print("Successfully generated project_overview.docx")

if __name__ == "__main__":
    create_document()
